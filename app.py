import re
import docx
import streamlit as st
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 页面配置
st.set_page_config(page_title="智能中文论文排版与降AI率助手", page_icon="📝", layout="wide")

st.title("📝 智能中文论文排版与数学公式乱码清洗助手")
st.markdown("上传您的中文 Word 论文 (.docx)，程序将自动过滤 AI 痕迹、规范中文宋体与西文排版、修复数学公式乱码。")

class ThesisProcessor:
    def __init__(self, doc_file):
        self.doc = docx.Document(doc_file)

    def clean_and_repair_text(self):
        """核心清洗与修复逻辑：针对中文文本及 LaTeX 符号"""
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            
            # 1. 清理复制产生的不可见乱码方块
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\ufffd]', '', text)
            
            # 2. 转换常见的 AI LaTeX 公式符号为标准中文可读符号
            text = text.replace(r"\times", "×").replace(r"\div", "÷")
            text = text.replace(r"\alpha", "α").replace(r"\beta", "β").replace(r"\gamma", "γ")
            text = text.replace("$", "")
            
            paragraph.text = text

    def apply_typography(self):
        """应用严格的中文论文排版规范（中文字体宋体，西文Times New Roman，1.5倍行距）"""
        for paragraph in self.doc.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.space_after = Pt(6)
            
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)  # 小四号
                
                # 精准设置中文字体兼容（确保中文显示为宋体）
                rPr = run._r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:eastAsia'), '宋体')
                rPr.append(rFonts)

    def save_file(self, output_path):
        self.doc.save(output_path)

# 文件上传组件
uploaded_file = st.file_uploader("请选择您的中文 Word 论文文件 (.docx)", type=["docx"])

if uploaded_file is not None:
    st.info(f"已成功加载文件: {uploaded_file.name}")
    
    if st.button("🚀 开始一键排版与公式清洗"):
        with st.spinner("正在处理中，请稍候..."):
            try:
                processor = ThesisProcessor(uploaded_file)
                processor.clean_and_repair_text()
                processor.apply_typography()
                
                output_filename = "formatted_" + uploaded_file.name
                processor.save_file(output_filename)
                
                st.success("处理完成！")
                with open(output_filename, "rb") as f:
                    st.download_button(
                        label="📥 点击下载排版后的中文论文",
                        data=f,
                        file_name=output_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"处理过程中出现错误：{e}")
